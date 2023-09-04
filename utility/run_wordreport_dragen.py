#!/usr/bin/env python
# -*- coding: utf-8 -*-
#shent2 on 2018/09/13: modify the value of trueseq (library prep protocol) and add "check before delivery"
#shent2 on 2018/09/27: remove "check before delivery" and make "library prep protocol" bold in the first table
#2019/11/18 shent2: adjust the report to fit metrics generated by dragen v3.4.9
#2020/10/29 shent2: retrieve the flowcell ID from RunInfo.xml
#2020/11/06 shent2: retrieve NovaSeq WorkflowType and FlowCellMode from RunParameters.xml
#2021/03/27 shent2: retrieve the reference from projectName/config.py. If projectName/config.py does not exist, the reference will be retrieved from LabQC metadata.
#2021/05/10 chenx3: added read length to tracking excel file
#2021/06/16 shent2: modified the path to access case-incensitive [rR]unParameters.xml and [rR]unInfo.xml using glob.glob
#2021/07/21 shent2: 1. added arguments of --analysisFolder and --projectName
#                   2. modified to retrieve the NAS ID from the project name rather than the LIMS Metadata.txt
#2021/08/26 shent2: 1. retrieved the read lengths from {analysis}/fastq if RunInfo.xml don't exist 
#                   2. updated the report templates for rna, chip, and nopipe and the formats in the Word report
#                   3. added the argument to assign the project name so that the NAS ID is parsed from the current project name instead of the one on LIMS
#                   4. parsed dragen.log to retrieve the Dragen version and the command to update the table of Software and Parameters
#2021/09/15 shent2: modified the Word report template to report the number of somatic variants in pairs (of case vs control) and tumor-only separately
#2021/09/28 shent2: removed raw coverage (rawcovmin, rawcovmax) from the Word report and modified the plot of Mapped Reads Statistics for mapped coverage
"""
Created on Tue Nov 21 23:58:30 2017

@author: Jack

uasge: python run_wordreport_dragen.py -i MikeBustin_CS022033_12Chiplib_110817_HVNYWBGX3.xlsx

"""
import os, sys, re, glob, gzip
import pandas as pd
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from xml.dom import minidom

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

import numpy as np

import warnings,argparse
parser = argparse.ArgumentParser(description="""Parse an Excel report to generate
                                                a Word report for delivery
                                                for data demultiplexing.""")
parser.add_argument("-i", "--input", metavar="report.xlsx", dest="finput",
                    action="store", type=str, required=True,
                    help="input report.xlsx (required)")
parser.add_argument("-r", "--run", metavar="runName", dest="runName",
                    action="store", type=str, required=False,
                    help="run name (optional)")
parser.add_argument("-p", "--projectName", metavar="projectName", dest="projectName",
                    action="store", type=str, required=False,
                    help="project name (optional)")
parser.add_argument("-af", "--analysisFolder", metavar="analysisFolder", dest="analysisFolder",
                    action="store", type=str, required=False,
                    help="analysis folder (optional)")
parser.add_argument("-o", "--output", metavar="report.docx", dest="foutput",
                    action="store", default="report.docx", type=str,
                    help="output file name (default: %(default)s)")
args = parser.parse_args()

def replace_string(doc, before, after):
    for p in doc.paragraphs:
        if before in p.text:
            inline = p.runs

            for i in range(len(inline)):
                if before in inline[i].text:
                    text = inline[i].text.replace(before, str(after))
                    inline[i].text = text
                    #return text
    return 1

def retrieveReadLength(fileName):
    header = ""
    readSeq = ""
    with gzip.open(fileName, 'rb') as IN:
        lineNum = 1
        for line in IN:
            if lineNum == 1:
                header = line[0:-1]
            elif lineNum == 2:
                readSeq = line[0:-1]
            else:
                break
            lineNum += 1
    return(str(len(readSeq)))

def get_xml_value(xmldoc, tag_name):
    elements = xmldoc.getElementsByTagName(tag_name)
    if elements:
        return elements[0].childNodes[0].nodeValue
    return "Unknown"

def parseXml(xmlRunParametersPath, xmlRunInfoPath, instrument, readLengths, indexLengths):
    xmldoc = minidom.parse(xmlRunParametersPath)

    # Determine RTA version based on instrument
    rta_tags = {
        "NovaSeq 6000": "RtaVersion",
        "NextSeq 2000": "RtaVersion",
        "iSeq": "RtaVersion",
        "HiSeq": "RTAVersion",
        "NextSeq 550": "RTAVersion",
        "MiSeq": "RTAVersion"
    }
    tagRTA = rta_tags.get(instrument, "")
    RTAVersion = get_xml_value(xmldoc, tagRTA) if tagRTA else "Unknown"
    if instrument == "NovaSeq Xplus":
        RTAVersion = "RTA4"

    # Extract other XML values
    workFlowTypeValue = get_xml_value(xmldoc, 'WorkflowType')
    flowcellModeValue = get_xml_value(xmldoc, 'FlowCellMode')
    chemistryValue = get_xml_value(xmldoc, 'Chemistry')
    chemistryVersionValue = get_xml_value(xmldoc, 'ChemistryVersion')

    # Special case for ConsumableInfo
    consumableInfo = xmldoc.getElementsByTagName('ConsumableInfo')
    if consumableInfo:
        flowcellModeValue = get_xml_value(consumableInfo[0], 'Mode')

    # Parse xmlRunInfoPath
    xmldoc = minidom.parse(xmlRunInfoPath)
    runInfoReads = xmldoc.getElementsByTagName('Read')
    for runInfoRead in runInfoReads:
        cycles = runInfoRead.getAttribute("NumCycles")
        if runInfoRead.getAttribute("IsIndexedRead") == "N":
            readLengths.append(cycles)
        else:
            indexLengths.append(cycles)

    flowcell = get_xml_value(xmldoc, 'Flowcell')

    return (RTAVersion, flowcell, workFlowTypeValue, flowcellModeValue, chemistryValue, chemistryVersionValue)

def getRunPath(sRunName):
    # Define a mapping for instruments and their respective paths
    instruments = {
        "A": ("NovaSeq 6000", ["RawData", "RawData_NovaSeq"]),
        "L": ("NovaSeq Xplus", ["RawData", "RawData_Xplus"]),
        "J": ("HiSeq", ["RawData", "RawData_HiSeq"]),
        "D": ("HiSeq", ["RawData", "RawData_HiSeq"]),
        "N": ("NextSeq 550", ["RawData_NextSeq", "RawData_NextSeq"]),
        "V": ("NextSeq 2000", ["RawData_NextSeq", "RawData_NextSeq"]),
        "M": ("MiSeq", ["RawData_MiSeq", "RawData_MiSeq"]),
        "F": ("iSeq", ["RawData_iSeq", "RawData_iSeq"])
    }

    base_paths = [
        '/is2/projects/CCR-SF/scratch/illumina',
        '/mnt/ccrsf-raw/illumina'
    ]

    instrument_key = sRunName.group(2)[0]
    instrument_name, paths = instruments.get(instrument_key, ('', ''))

    # Loop through the base paths and the instrument paths to find the correct directory
    for base in base_paths:
        for path in paths:
            full_path = os.path.join(base, path, sRunName.group(0))
            if os.path.isdir(full_path):
                return instrument_name, full_path

    return instrument_name, ''

def get_xml_paths(runPath, runName):
    xmlRunParametersPath = glob.glob(os.path.join(runPath, runName, "[rR]unParameters.xml"))[0]
    xmlRunInfoPath = glob.glob(os.path.join(runPath, runName, "[rR]unInfo.xml"))[0]
    return xmlRunParametersPath, xmlRunInfoPath

def get_fastq_path(args, projectName):
    if args.analysisFolder:
        return f'{args.analysisFolder}/fastq'
    else:
        return f'{projectName}/fastq'

def retrieve_read_lengths_from_fastq(fastqPath):
    oFileNameR1 = re.compile('(.+)_R1_001.fastq.gz')
    oFileNameR2 = re.compile('(.+)_R2_001.fastq.gz')
    fileNameR1 = ''
    fileNameR2 = ''
    flagR1, flagR2 = 0, 0
    for fileName in os.listdir(fastqPath):
        sFileNameR1 = oFileNameR1.search(fileName)
        sFileNameR2 = oFileNameR2.search(fileName)
        if sFileNameR1:
            flagR1 = 1
            fileNameR1 = fileName
        elif sFileNameR2:
            flagR2 = 1
            fileNameR2 = fileName

    readLengths = []
    if flagR1 == 1 and flagR2 == 1:
        readLengths.append(retrieveReadLength(f'{fastqPath}/{fileNameR1}'))
        readLengths.append(retrieveReadLength(f'{fastqPath}/{fileNameR2}'))
    elif flagR1 == 1 and flagR2 == 0:
        readLengths.append(retrieveReadLength(f'{fastqPath}/{fileNameR1}'))
    else:
        sys.stdout.write(f'The fastq file name in {fastqPath} is not standard, so the read length could not be retrieved.\n')
    return readLengths

def process_run_name(sRunName, args, projectName, readLengths, indexLengths):
    (instrument, runPath) = getRunPath(sRunName)
    if os.path.isdir(os.path.join(runPath, runName)):
        xmlRunParametersPath, xmlRunInfoPath = get_xml_paths(runPath, runName)
        (RTAVersion, flowcell, workFlowType, flowcellMode, chemistry, chemistryVersion) = parseXml(xmlRunParametersPath, xmlRunInfoPath, instrument, readLengths, indexLengths)
        if flowcellMode == 'NextSeq 2000 P3 Flow Cell Cartridge':
            flowcellMode = 'P3'
        elif flowcellMode == 'NextSeq 1000/2000 P2 Flow Cell Cartridge':
            flowcellMode = 'P2'
    else:
        sys.stdout.write(f'{xmlRunParametersPath}/{xmlRunInfoPath} does not exist.\nRetriving read lengths from fastq.gz...\n')
        fastqPath = get_fastq_path(args, projectName)
        readLengths = retrieve_read_lengths_from_fastq(fastqPath)
    return RTAVersion

def get_sheet_data(data, sheet_name):
    if sheet_name in data.sheet_names:
        return data.parse(sheet_name)
    return None

def get_somatic_data(data):
    somatic_data = get_sheet_data(data, 'Somatic')
    tumor_only_data = get_sheet_data(data, 'TumorOnly')
    
    if somatic_data and not tumor_only_data:
        return somatic_data, somatic_data.shape[0], 0, 1
    elif not somatic_data and tumor_only_data:
        return tumor_only_data, 0, tumor_only_data.shape[0], 1
    elif somatic_data and tumor_only_data:
        combined_data = pd.concat([somatic_data, tumor_only_data], ignore_index=True, sort=False)
        return combined_data, somatic_data.shape[0], tumor_only_data.shape[0], 2
    else:
        sys.stdout.write("No somatic variant calling spreadsheet.\n")
        return None, 0, 0, 0

def get_project_name_and_csac(args, lqc):
    oProjectName = re.compile("([-a-zA-Z]+)_([CS0-9]+)_(\w+)_(\d+)")
    if args.projectName:
        projectName = args.projectName
    else:
        projectName = lqc[lqc.columns[3]][0]
        sys.stdout.write(f'The project name is not assigned in the command line.\nThe NAS of {csac} is retrieved from LIMS Metadata.txt.\n')

    sProjectName = oProjectName.search(projectName)
    csac_from_lims = str(lqc[lqc.columns[10]][0])

    if sProjectName:
        csac_from_project = sProjectName.group(2)
        if csac_from_lims != csac_from_project:
            sys.stdout.write(
                f'{csac_from_lims} in LIMS Metadata.txt is inconsistent with the NAS in the project name, {projectName}.\n'
                f'The current NAS is set as {csac_from_project}.\n'
            )
            return projectName, csac_from_project
    else:
        sys.stdout.write(f'Cannot recognize the NAS from {projectName} in LIMS Metadata.txt\n{csac_from_lims} is retrieved from LIMS Metadata.txt.\n')
    
    return projectName, csac_from_lims

def parse_data(args):
    excelfile = args.finput
    data = pd.ExcelFile(excelfile)
    
    lqc = get_sheet_data(data, 'LabQC')
    summary = get_sheet_data(data, 'Summary')
    vt, numPairedCases, numTumorOnlyCases, vtSomaticFlag = get_somatic_data(data)
    vtGermline = get_sheet_data(data, 'JointGenotyping')
    vtGermlineFlag = 1 if vtGermline else 0

    if "EXOME" in lqc['Application'][0].upper():
        docfile = '/mnt/ccrsf-ifx/Software/scripts/lib/wordreporttemp/dragen_exome.docx'
        seqType = "Exome"
    else:
        docfile = '/mnt/ccrsf-ifx/Software/scripts/lib/wordreporttemp/dragen_wgs.docx'
        seqType = lqc['Application'][0]

    doc = Document(docfile)
    PIName, labContact, bioinfoContact = lqc[lqc.columns[1]][0], lqc[lqc.columns[2]][0], lqc[lqc.columns[4]][0]
    projectName, csac = get_project_name_and_csac(args, lqc)
    stip, sitr = str(len(lqc)), str(len(summary))
    date = time.strftime("%x")

    return doc, PIName, labContact, bioinfoContact, projectName, csac, stip, sitr, date

def get_germline_data(data):
    if "JointGenotyping" in data.sheet_names:
        return get_sheet_data(data, 'JointGenotyping'), 1
    else:
        sys.stdout.write("No germline joint genotyping spreadsheet.\n")
        return None, 0

def get_document_template(lqc):
    if "EXOME" in lqc['Application'][0].upper():
        return '/mnt/ccrsf-ifx/Software/scripts/lib/wordreporttemp/dragen_exome.docx', "Exome"
    else:
        return '/mnt/ccrsf-ifx/Software/scripts/lib/wordreporttemp/dragen_wgs.docx', lqc['Application'][0]

def plot_total_and_mapped_reads(summary):
    x = list(summary[summary.columns[0]])
    y0 = list(summary['Total input reads'])
    y1 = list(summary['% Total mapped reads'])
    x_pos = np.arange(len(x))

    fig, ax1 = plt.subplots()
    ax1.bar(x_pos, y0, align='center', alpha=0.5, label='Total Reads')
    ax1.set_ylabel('Total Reads')
    ax1.set_xticks(x_pos)
    ax1.set_xticklabels(x, rotation=45, ha='right')
    ax1.set_title('Total Reads & Percent of Mapped Reads', fontsize=30)
    ax1.legend(loc=3)

    ax2 = ax1.twinx()
    ax2.plot(x_pos, y1, 'r-o', label='Percent of Mapped Reads')
    ax2.set_ylabel('Percent of Mapped Reads', color='r')
    ax2.set_ylim(0, 100)
    ax2.legend(loc=4)

    fig.set_size_inches(12, 8)
    fig.savefig('matplt1.png', dpi=300, bbox_inches='tight')
    plt.show()
    plt.clf()

def plot_mapped_reads_statistics(summary, lqc):
    x = list(summary[summary.columns[0]])
    y0 = list(summary['% Uniquely mapped reads'])
    x_pos = np.arange(len(x))

    if "EXOME" in lqc['Application'][0].upper():
        y1 = list(summary['% of reads on target region'])
        lb = "Percent Reads Mapped On Target"
        y2 = list(summary['% of target region with coverage above 20x'])
        lbb = "% of target region with coverage above 20x"
    else:
        y1 = list(summary['Mean of mapped coverage'])
        lb = "Mean of mapped coverage"
        y2 = list(summary['% of genome with coverage above 20x'])
        lbb = "% of genome with coverage above 20x"

    plt.plot(x_pos, y0, 'b-o', label='Percent of Uniquely mapped reads')
    plt.plot(x_pos, y1, 'r-o', label=lb)
    plt.plot(x_pos, y2, 'g-o', label=lbb)
    plt.xticks(x_pos, x, rotation=45, ha='right')
    plt.ylabel('Percent')
    plt.title('Mapped Reads Statistics', fontsize=30)
    plt.ylim(0, 100)
    plt.legend(loc=4)

    fig = plt.gcf()
    fig.set_size_inches(12, 8)
    fig.savefig('matplt2.png', dpi=300, bbox_inches='tight')
    plt.clf()

def add_plots_to_doc(doc, filenm):
    p = doc.tables[3].rows[0].cells[0].add_paragraph()
    r = p.add_run()
    r.add_picture('matplt1.png', height=Inches(4))

    p = doc.tables[3].rows[1].cells[0].add_paragraph()
    r = p.add_run()
    r.add_picture('matplt2.png', height=Inches(4))

    doc.save(filenm)

def replace_string(doc, placeholder, value):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(placeholder, value)

def populate_table(cell, df, title):
    cell.add_paragraph(text=f"\n\n{title}:\n")
    rows, columns = df.shape
    tb = cell.add_table(1, columns)
    tb.style = 'Table Grid'
    h = list(df)
    head = tb.rows[0].cells
    for idx, name in enumerate(h):
        paragraph = head[idx].paragraphs[0]
        run = paragraph.add_run(name)
        run.bold = True
    for r in range(rows):
        cells = tb.add_row().cells
        for col in range(columns):
            cells[col].text = str(df.iat[r, col])

def handle_email_message(docfile, lqc, PIName, rc, rc1, rcGermline, flowcell, st, csac, instrument, summary, vtSomaticFlag, vt, vtPaired, vtTumorOnly, vtGermlineFlag, vtGermline, emailfile):
    doc = Document(docfile)
    app = str(lqc['Application'][0])
    if '_' in PIName:
        pilast = PIName.replace("_", " ").split()[1]
    elif ',' in PIName:
        pilast = PIName.split(', ')[0]
    else:
        pilast = PIName
    ttext = f"Mapping:\n\n{rc.text}\n\nVariant Calling:\n\n{rc1.text}\n\n{rcGermline.text}"

    replace_string(doc, "{App}", app)
    replace_string(doc, "{Flowcell1}", flowcell)
    replace_string(doc, "Pilast", pilast)
    replace_string(doc, "{ST}", st)
    replace_string(doc, "{CSAC}", csac)
    replace_string(doc, "{Flowcell2}", flowcell)
    replace_string(doc, "{MT}", instrument)

    doc.paragraphs[14].text = ttext

    populate_table(doc.tables[0].rows[0].cells[0], summary, "Mapping")

    if vtSomaticFlag == 1:
        populate_table(doc.tables[0].rows[0].cells[0], vt, "Somatic Variant Calling")
    elif vtSomaticFlag == 2:
        populate_table(doc.tables[0].rows[0].cells[0], vtPaired, "Somatic Variant Calling (Paired)")
        populate_table(doc.tables[0].rows[0].cells[0], vtTumorOnly, "Somatic Variant Calling (Tumor-only)")

    if vtGermlineFlag == 1:
        populate_table(doc.tables[0].rows[0].cells[0], vtGermline, "Germline Joint Genotyping")

    doc.save(emailfile)

def construct_record(excelfile, summary, lqc, readLengths):
    out = ""
    out += os.getcwd().split("/")[-1] + ","
    out += os.path.basename(excelfile).split('.')[0] + ","
    adate = time.strftime('%m/%d/%y', time.gmtime(os.path.getctime(excelfile)))
    out += str(adate) + ",cronjob,," + str(adate) + ","
    gb = str(float(sum(summary[summary.columns[1]]))/1000)
    out += gb + ","
    out += str(len(summary)) + ","
    out += str(lqc['Application'][0]) + ","
    if len(readLengths) == 1:
        out += readLengths[0] + "__0,"
    elif len(readLengths) == 2:
        out += readLengths[0] + "__" + readLengths[1] + ","
    else: 
        out += "Customized,"
    out += str(adate) + "\n"
    return out

def record_exists(csvrecord, excelfile):
    with open(csvrecord, 'r') as ck:
        ckk = ck.read()
    return os.path.basename(excelfile).split('.')[0] in ckk

def update_record(csvrecord, excelfile, summary, lqc, readLengths):
    if not record_exists(csvrecord, excelfile):
        out = construct_record(excelfile, summary, lqc, readLengths)
        with open(csvrecord, 'a') as cr:
            cr.write(out)


def main(argv):
    path = os.getcwd()
    paths = path.split("/")
    if args.runName is None:
        runName = paths[-1]
    else:
        runName = args.runName
    instrument = "Unknown"
    #rawDataPath = "/is2/projects/CCR-SF/scratch/illumina"
    runPath = ''
    xmlRunParametersPath = ""
    xmlRunInfoPath = ""
    readLengths = []
    indexLengths = []
    RTAVersion = "Unknown"
    flowcell = "Unknown"
    workFlowType = "Unknown"
    flowcellMode = "Unknown"
    chemistry = "Unknown"
    chemistryVersion = "Unknown"
    oRunName = re.compile("(\d{6,})_([A-Z0-9]+)_(\d+)_([-A-Z0-9]+)")
    sRunName = oRunName.search(runName)

    if sRunName:
        RTAVersion = process_run_name(sRunName, args, projectName, readLengths, indexLengths)
    else:
        sys.stdout.write("The run name is not recognized from the current path.\n")
        sys.stdout.write("Please assign the run name using,\n")
        sys.stdout.write("run_wordreport_dragen.py -i report.xlsx -r run_name\n")
        sys.stdout.write("The current Word report is generated without checking RunInfo.xml and RunParameter.xml\n")
        RTAVersion = ""

    excelfile = args.finput
    filenm, emailfile = get_file_names(excelfile)
    data = get_data_from_excel(excelfile)
    
    lqc = get_sheet_data(data, 'LabQC')
    summary = get_sheet_data(data, 'Summary')
    
    vt, vtSomaticFlag = get_somatic_data(data)
    vtGermline, vtGermlineFlag = get_germline_data(data)
    
    docfile, seqType = get_document_template(lqc)
    doc = Document(docfile)



    plot_total_and_mapped_reads(summary)
    plot_mapped_reads_statistics(summary, lqc)
    add_plots_to_doc(doc, 'output_filename.docx')

    update_record(csvrecord, excelfile, summary, lqc, readLengths)



if __name__ == "__main__":
    main(sys.argv[1:])
